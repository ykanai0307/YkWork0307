import MeCab
import csv
import sys
from .dirPath import dirPathClass
from os.path import abspath, join, split
from docx import Document

class tokenizerMainClass:

    #コンストラクタ
    def __init__(self):
        self._file = ""
        self._dirPathIns = dirPathClass()
        # Workディレクトリ指定
        self.self_dir = 'C:\\YkWork\\pythonWork\\'
        self.in_dir = 'C:\\YkWork\\pythonWork\\In\\'
        # mainProject直下のstaticを参照
        self.out_dir = abspath(join(__file__, '../../../static/result')) + '/'
        self.input_file = 'input.txt'
        # 出力CSV
        self.csv_file = 'result.csv'

    # 入力されるファイル一覧取得(関数)
    def InputFileList(self):
        row_list = ""
        try:
            input_file_list = open(self.in_dir + self.input_file, 'r', encoding = "utf_8")
            if not input_file_list is None:
                read_row = input_file_list.read()
                row_list = read_row.split('|*|')
            input_file_list.close()
        except Exception as e:
            raise e
        return row_list

    # 単語頻出度カウント(関数)
    def WordFrequencyCount(self,word,wordFreq_dic):
        try:
            if word in wordFreq_dic:
                wordFreq_dic[word] +=1
            else:
                wordFreq_dic.setdefault(word, 1)
        except Exception as e:
            raise e
        return wordFreq_dic

    # 単語分解(関数)
    def TextExtractWord(self,text,wordFreq_dic):
        try:
            # 特定の品詞の単語を抽出
            mecab = MeCab.Tagger()
            mecab.parse('')
            node="" #node = mecab.parseToNode(text)
            
            while node:
                if node.feature.split(",")[0] == "名詞":
                    word = node.surface
                    wordFreq_dic = self.WordFrequencyCount(word,wordFreq_dic)
                elif node.feature.split(",")[0] =="動詞":
                    word = node.surface
                    wordFreq_dic = self.WordFrequencyCount(word,wordFreq_dic)
                elif node.feature.split(",")[0] == "形容詞":
                    word = node.surface
                    wordFreq_dic = self.WordFrequencyCount(word,wordFreq_dic)
                elif node.feature.split(",")[0] == "形容動詞":
                    word = node.surface
                    wordFreq_dic = self.WordFrequencyCount(word,wordFreq_dic)
                else:pass
                node = node.next
        except Exception as e:
            raise e
        return wordFreq_dic

    # 降順に並び替え(関数)
    def SortList(self,wordcount_output):
        wordcount_output = sorted(wordcount_output, key = lambda x:x[1], reverse=True)
        return wordcount_output

    # 辞書リストを取り出す(関数)
    def DictGetItem(self,wordFreq_dic):
        wordcount_output = []
        try:
            for item in wordFreq_dic.items():
                wordcount_output.append(item)
        except Exception as e:
            raise e
        return wordcount_output

    # CSV出力(関数)
    def CsvOut(self,wordcount_output):
        with open(self.out_dir + self.csv_file, "w", encoding="utf-8") as f:
            writer = csv.writer(f, lineterminator="\n")
            writer.writerows(wordcount_output)

    # 解析テキスト取得(単一ファイル対象)
    def MainToOneFile(self,input_file):
        result = ""
        wordFreq_dic = {}
        try:
            # ドキュメント開く
            doc = Document( self._dirPathIns.GetSelfToStaticUploadDirPath() + input_file)
            if not doc is None:
                text = ''
                for para in doc.paragraphs:
                    text += para.text
            wordFreq_dic = self.TextExtractWord(text,wordFreq_dic)
            wordcount_output = self.DictGetItem(wordFreq_dic)
            wordcount_output = self.SortList(wordcount_output)
            # CSV出力
            self.CsvOut(wordcount_output)
        except Exception as e:
            print(e)
            result = e
        return result
        
    # 解析テキスト取得(複数ファイル対象)
    def MainToManyFile(self,input_file_list):
        result = ""
        wordFreq_dic = {}
        try:
            for in_file in input_file_list:
                # ドキュメント開く
                doc = Document( self._dirPathIns.GetSelfToStaticUploadDirPath() + in_file)
                if not doc is None:
                    text = ''
                    for para in doc.paragraphs:
                        text += para.text
                wordFreq_dic = self.TextExtractWord(text,wordFreq_dic)
                wordcount_output = self.DictGetItem(wordFreq_dic)
                wordcount_output = self.SortList(wordcount_output)
            # CSV出力
            self.CsvOut(wordcount_output)
        except Exception as e:
            print(e)
            result = e
        return result
     
    def GetFileProperty(self):
        return self._file #getter
     
    def SetFileProperty(self, value):
        self._file = value #setter